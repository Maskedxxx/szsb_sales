from docx import Document
from docx.shared import Pt

# Создание нового документа
doc = Document()

# Заголовок
doc.add_heading("Сравнительная таблица LLM для российского рынка", level=1)

# Добавление требований
doc.add_paragraph("**Требования к LLM:**\n"
                  "- Российская\n"
                  "- Инфраструктурная локализация серверов (Sber, VK Cloud)\n"
                  "- Защищённый контур SaaS (облака РФ: СберCloud, VK Cloud)")

# Создание таблицы
table = doc.add_table(rows=1, cols=6)
table.style = "Light Shading Accent 1"

# Заголовки
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Провайдер / Платформа"
hdr_cells[1].text = "Модель"
hdr_cells[2].text = "Стоимость (за 1 млн токенов)"
hdr_cells[3].text = "Контекстное окно"
hdr_cells[4].text = "Поддержка LangChain"
hdr_cells[5].text = "Примечания (Эффективность / Отзывы)"

# Данные из анализа
data = [
    ["Yandex Cloud", "YandexGPT Lite", "Синхр.: $1.667\nАсинхр.: $0.834", "~8k", "Да (через OpenAI совместимость)", "Высокое качество на русском"],
    ["Yandex Cloud", "YandexGPT Pro", "Синхр.: $10.002\nАсинхр.: $5.001", "~8k", "Да", "Снижен уровень галлюцинаций до 6%"],
    ["SberCloud", "GigaChat Lite", "≈200 ₽", "До 200 стр.", "Да (SDK)", "Базовая версия"],
    ["SberCloud", "GigaChat Pro", "≈1500 ₽ / 750 ₽", "До 200 стр.", "Да", "Баланс цены и качества"],
    ["SberCloud", "GigaChat Max", "≈1950 ₽ / 970 ₽", "До 200 стр.", "Да", "Лучшая модель для русского языка"],
    ["Yandex Cloud", "Llama 3 70B", "Синхр.: $10.002\nАсинхр.: $5.001", "32k+", "Да", "Open-source через API"],
    ["Yandex Cloud", "Qwen 32B", "Синхр.: $1.667\nАсинхр.: $0.834", "32k+", "Да", "Open-source через API"],
    ["Self-Hosted", "Qwen 32B (~30B)", "≈$10 (эффективная)", "32k+", "Да", "Нужна GPU 24 ГБ VRAM"],
    ["Self-Hosted", "Llama 70B", "≈$20 (эффективная)", "32k+", "Да", "Требует ≥48 ГБ VRAM"],
]

# Заполнение таблицы
for row in data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val

# Сохранение документа
output_path = "/Users/mask/Documents/Проеты_2025/szsb_sales/Сравнительная_таблица_LLM.docx"
doc.save(output_path)
output_path
